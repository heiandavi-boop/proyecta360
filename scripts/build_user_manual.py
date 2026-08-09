from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "manual"
ASSET_DIR = OUT_DIR / "assets"
CROP_DIR = OUT_DIR / "cropped"
OUTPUT = OUT_DIR / "manual_usuario_proyecta360.docx"


SCREENSHOTS = {
    "login": ("00-login.png", "Pantalla de ingreso protegido"),
    "portfolio": ("01-portafolio.png", "Portafolio ejecutivo con PHS y senales de gestion"),
    "master": ("02-plan-maestro.png", "Plan Maestro / Gantt con WBS, ruta critica y tareas resumen"),
    "scrum": ("03-scrum.png", "Tablero Scrum con historias, estados y burndown"),
    "resources": ("04-recursos.png", "Recursos y capacidad"),
    "risks": ("05-riesgos.png", "Riesgos del proyecto"),
    "conversations": ("06-conversaciones.png", "Conversaciones, decisiones y bloqueos"),
    "knowledge": ("07-conocimiento.png", "Conocimiento, entregables y evidencias"),
    "ai": ("08-ia.png", "IA del Proyecto: hallazgos y recomendaciones"),
    "budget": ("09-presupuesto.png", "Presupuesto mensual por rubro"),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, True)
        set_cell_shading(hdr[i], "EAF1FF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    r.font.size = Pt(10)
    p.add_run(f"\n{body}").font.size = Pt(9)
    doc.add_paragraph()


def crop_images() -> dict[str, Path]:
    CROP_DIR.mkdir(parents=True, exist_ok=True)
    result: dict[str, Path] = {}
    for key, (filename, _) in SCREENSHOTS.items():
        src = ASSET_DIR / filename
        dst = CROP_DIR / filename
        with Image.open(src) as image:
            width, height = image.size
            crop_height = min(height, 1080)
            cropped = image.crop((0, 0, width, crop_height))
            cropped.save(dst)
        result[key] = dst
    return result


def add_screenshot(doc: Document, path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.4))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(8)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for name, size, color in [
        ("Title", 24, "0F172A"),
        ("Heading 1", 17, "0F172A"),
        ("Heading 2", 13, "1D4ED8"),
        ("Heading 3", 11, "334155"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual de Usuario Proyecta360")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(15, 23, 42)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Guia completa de uso de la aplicacion")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(71, 85, 105)
    doc.add_paragraph()
    add_note(
        doc,
        "Alcance del manual",
        "Este documento cubre los modulos visibles de la aplicacion, los flujos operativos, las funciones transversales, los roles, las importaciones, exportaciones, evidencias, IA y recomendaciones.",
    )
    meta = [
        ["Aplicacion", "Proyecta360"],
        ["Version documentada", "MVP actualizado con PHS ejecutivo"],
        ["Fecha", date.today().isoformat()],
        ["Usuario de referencia para capturas", "alejandra@proyecta360.ai - Project Manager"],
    ]
    add_table(doc, ["Campo", "Detalle"], meta)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("Contenido", level=1)
    rows = [
        ["1", "Vision general"],
        ["2", "Ingreso, idioma y roles"],
        ["3", "Elementos comunes de la interfaz"],
        ["4", "Portafolio"],
        ["5", "Creacion y edicion de proyectos"],
        ["6", "Plan Maestro / Gantt"],
        ["7", "Scrum"],
        ["8", "Recursos"],
        ["9", "Presupuesto"],
        ["10", "Riesgos"],
        ["11", "Conversaciones"],
        ["12", "Conocimiento, entregables y evidencias"],
        ["13", "IA del Proyecto"],
        ["14", "Importacion, exportacion y contratos"],
        ["15", "Flujos recomendados"],
        ["16", "Glosario y criterios de salud"],
    ]
    add_table(doc, ["#", "Seccion"], rows)
    doc.add_page_break()


def add_overview(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("1. Vision general", level=1)
    doc.add_paragraph(
        "Proyecta360 es una plataforma para gestionar proyectos con un enfoque hibrido: combina control tradicional tipo PMP, ejecucion agil tipo Scrum/Kanban, trazabilidad documental, conversaciones operativas e inteligencia artificial."
    )
    add_bullets(doc, [
        "Centraliza portafolio, cronograma, riesgos, recursos, Scrum, conversaciones, evidencias e IA.",
        "Calcula senales ejecutivas como avance real vs esperado, desviacion presupuestal y Project Health Score.",
        "Permite crear proyectos manualmente o importarlos desde CSV con entidades relacionadas.",
        "Mantiene permisos por rol y exige aprobacion humana antes de aplicar recomendaciones de IA.",
    ])
    add_screenshot(doc, shots["portfolio"], SCREENSHOTS["portfolio"][1])


def add_auth(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("2. Ingreso, idioma y roles", level=1)
    add_screenshot(doc, shots["login"], SCREENSHOTS["login"][1])
    doc.add_heading("Ingreso", level=2)
    add_numbered(doc, [
        "Abrir la aplicacion en el navegador.",
        "Ingresar correo y contrasena autorizados.",
        "Presionar Ingresar.",
        "Al autenticarse, la aplicacion carga el portafolio y el proyecto activo.",
    ])
    doc.add_heading("Usuarios iniciales", level=2)
    add_table(doc, ["Rol", "Correo", "Permisos"], [
        ["Administrador", "admin@proyecta360.local", "Lectura, escritura y administracion de configuraciones sensibles."],
        ["Project Manager", "alejandra@proyecta360.ai", "Lectura y escritura operativa sobre proyectos, tareas, riesgos, recursos, conversaciones, evidencias e IA."],
        ["Consulta", "consulta@proyecta360.local", "Lectura. No puede crear, editar, aplicar recomendaciones ni cargar archivos."],
    ])
    doc.add_heading("Selector de idioma", level=2)
    doc.add_paragraph("La barra superior incluye botones de idioma. Permite cambiar los textos de la interfaz entre los catalogos disponibles: ingles, chino, hindi, espanol y arabe.")


def add_common(doc: Document) -> None:
    doc.add_heading("3. Elementos comunes de la interfaz", level=1)
    add_bullets(doc, [
        "Barra lateral: acceso a Portafolio, Plan Maestro, Scrum, Recursos, Presupuesto, Riesgos, Conversar, Conocimiento e IA Proyecto.",
        "Encabezado del proyecto: muestra nombre, descripcion y vista activa.",
        "KPIs del proyecto: aparecen en las vistas de gestion y resumen avance general, fecha fin calculada, riesgos abiertos, ruta critica y presupuesto ejecutado.",
        "Botones primarios: crean registros o ejecutan acciones principales.",
        "Tablas: permiten revisar registros existentes, estados y responsables.",
        "Permisos: si el usuario tiene rol Consulta, los controles de escritura se ocultan o quedan deshabilitados.",
    ])
    add_note(doc, "Regla operativa", "La aplicacion no cambia datos criticos sin una accion explicita del usuario. Las recomendaciones de IA quedan pendientes hasta que alguien las apruebe o aplique.")


def add_portfolio(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("4. Portafolio", level=1)
    add_screenshot(doc, shots["portfolio"], SCREENSHOTS["portfolio"][1])
    doc.add_paragraph("El Portafolio es la vista ejecutiva para revisar todos los proyectos y decidir donde actuar primero.")
    add_table(doc, ["Elemento", "Funcion"], [
        ["Buscador", "Filtra proyectos por nombre, PM o metodologia."],
        ["Filtro de estado", "Permite aislar proyectos saludables, en riesgo o criticos."],
        ["PHS", "Project Health Score ponderado por cronograma, presupuesto y riesgos."],
        ["Avance real / esperado", "Compara avance ejecutado contra el avance que deberia existir a la fecha."],
        ["Presupuesto ejecutado / esperado", "Muestra si el gasto acumulado va por encima o por debajo de lo esperado."],
        ["Hitos en riesgo", "Cuenta hitos vencidos o comprometidos y tareas criticas abiertas."],
        ["Proximo hito", "Muestra el hito futuro mas cercano para anticipar seguimiento."],
        ["Abrir", "Entra al proyecto seleccionado y abre su Plan Maestro."],
        ["Editar", "Permite modificar la ficha del proyecto."],
    ])
    doc.add_heading("Como usarlo", level=2)
    add_numbered(doc, [
        "Revisar primero proyectos con estado Critico o En riesgo.",
        "Mirar el PHS y sus componentes C, P y R para saber si el problema esta en cronograma, presupuesto o riesgos.",
        "Comparar avance real vs esperado para identificar retrasos reales.",
        "Comparar presupuesto ejecutado vs esperado para detectar sobreejecucion.",
        "Abrir el proyecto que requiera intervencion.",
    ])


def add_project_create(doc: Document) -> None:
    doc.add_heading("5. Creacion y edicion de proyectos", level=1)
    doc.add_paragraph("Desde Portafolio, el boton Nuevo proyecto abre un formulario por pestanas. La creacion es deliberadamente ligera: Problema, Alcance y Contexto enriquecen la inteligencia del proyecto, pero no bloquean la creacion.")
    add_table(doc, ["Pestana", "Campos / uso"], [
        ["General", "Nombre, codigo, Project Manager, sponsor, area solicitante, tipo, metodologia y descripcion."],
        ["Problema", "Problema o brecha, situacion actual y consecuencia de no ejecutar el proyecto."],
        ["Alcance", "Objetivo general, objetivos especificos, indicadores, alcance incluido, fuera de alcance, criterios de exito, supuestos y restricciones."],
        ["Contexto", "Contexto general, entorno politico, geografico, socioeconomico, cultural, institucional, partes interesadas, dependencias externas y restricciones regulatorias."],
        ["Planificacion", "Fecha de inicio, fecha de finalizacion prevista, moneda y presupuesto total."],
    ])
    add_bullets(doc, [
        "Prioridad y estado inicial no se piden en el formulario porque el sistema debe derivar senales reales desde datos operativos.",
        "Equipo responsable se gestiona mejor desde Recursos.",
        "Los campos de problema, alcance y contexto alimentan IA, recomendaciones, riesgos e indicadores cuando estan disponibles.",
    ])
    doc.add_heading("Importar desde CSV", level=2)
    doc.add_paragraph("El panel de importacion permite cargar un archivo CSV con columna entity. Puede incluir project, component, resource, budget, task, dependency, sprint, story, risk, deliverable, conversation_thread y conversation_message.")


def add_master(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("6. Plan Maestro / Gantt", level=1)
    add_screenshot(doc, shots["master"], SCREENSHOTS["master"][1])
    doc.add_paragraph("El Plan Maestro organiza el proyecto como una WBS: fases, actividades, tareas e hitos. La jerarquia se maneja con sangria; las dependencias se manejan con Vincular.")
    add_table(doc, ["Funcion", "Descripcion"], [
        ["+ Tarea", "Crea una actividad ejecutable con inicio, duracion, responsable, avance y tipo."],
        ["+ Hito", "Crea un punto de control sin duracion."],
        ["Indentar / Desindentar", "Cambia la jerarquia WBS. Una tarea con hijos se vuelve tarea resumen."],
        ["Agregar subtarea", "Crea una tarea hija dentro de la tarea seleccionada."],
        ["Agregar tarea debajo", "Inserta una tarea en el mismo nivel."],
        ["Vincular", "Crea una dependencia de precedencia entre tareas. No cambia la jerarquia."],
        ["Expandir / Contraer", "Muestra u oculta tareas hijas en la WBS."],
        ["Ver ruta critica", "Filtra tareas que afectan la fecha final del proyecto."],
        ["Recalcular cronograma", "Actualiza fechas, duraciones, ruta critica y tareas resumen."],
        ["Sincronizar avance desde Scrum", "Actualiza avance de una tarea a partir de historias Scrum vinculadas."],
    ])
    add_note(doc, "Tareas resumen", "Cuando una tarea tiene hijas, sus fechas, duracion y avance se calculan desde sus descendientes. El usuario no debe editar manualmente esos valores.")


def add_scrum(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("7. Scrum", level=1)
    add_screenshot(doc, shots["scrum"], SCREENSHOTS["scrum"][1])
    doc.add_paragraph("El modulo Scrum permite ejecutar trabajo agil sin perder conexion con el Plan Maestro.")
    add_table(doc, ["Elemento", "Uso"], [
        ["Nueva historia", "Registra titulo, actividad del Plan Maestro, sprint, estado, puntos, responsable y prioridad."],
        ["Estados del tablero", "Incluye Por hacer, En progreso y Hecho; tambien permite crear estados personalizados."],
        ["Arrastrar historias", "Mueve historias entre columnas para actualizar su estado."],
        ["Reordenar columnas", "Permite cambiar el orden visual de estados del tablero."],
        ["Burndown", "Muestra progreso de puntos completados frente al total del sprint."],
        ["Vinculo con Plan Maestro", "Cada historia puede relacionarse con una tarea para alimentar avance y trazabilidad."],
    ])


def add_resources(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("8. Recursos", level=1)
    add_screenshot(doc, shots["resources"], SCREENSHOTS["resources"][1])
    doc.add_paragraph("Recursos centraliza las personas o capacidades asignadas al proyecto.")
    add_bullets(doc, [
        "Registrar nombre, rol, correo y capacidad disponible.",
        "Visualizar capacidad como porcentaje.",
        "Detectar disponibilidad o posible sobrecarga.",
        "Reutilizar responsables en tareas, historias, riesgos, componentes y entregables.",
    ])


def add_budget(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("9. Presupuesto", level=1)
    add_screenshot(doc, shots["budget"], SCREENSHOTS["budget"][1])
    doc.add_paragraph("Presupuesto registra el plan mensual por rubro y la ejecucion real. Cuando existe este plan, el PHS usa esos valores en lugar de estimar el gasto desde las tareas.")
    add_table(doc, ["Elemento", "Uso"], [
        ["Mes", "Periodo de planeacion en formato AAAA-MM."],
        ["Rubro", "Categoria presupuestal como Equipo, Servicios, Licencias, Viajes u otra definida por el equipo."],
        ["Planificado", "Monto esperado para el mes y rubro."],
        ["Ejecutado", "Monto real ejecutado o comprometido."],
        ["Notas", "Contexto financiero, soporte o explicacion de variaciones."],
        ["Indicadores", "Resumen de planificado total, ejecutado total, desviacion y fuente del PHS."],
        ["Tabla por rubro", "Acumula plan y ejecucion para comparar categorias."],
        ["Registros mensuales", "Lista cada entrada creada y permite eliminar registros incorrectos."],
    ])
    doc.add_heading("Como usarlo", level=2)
    add_numbered(doc, [
        "Abrir Presupuesto desde la barra lateral.",
        "Seleccionar el mes que se va a planear o actualizar.",
        "Registrar rubro, monto planificado, monto ejecutado y notas.",
        "Guardar la entrada.",
        "Revisar la desviacion. Si el ejecutado supera el plan acumulado, el componente presupuestal del PHS baja.",
    ])
    add_note(doc, "Relacion con PHS", "Si el proyecto no tiene registros presupuestales mensuales, la aplicacion conserva el calculo estimado desde tareas. En cuanto se registra el plan mensual, el PHS cambia a fuente plan mensual.")


def add_risks(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("10. Riesgos", level=1)
    add_screenshot(doc, shots["risks"], SCREENSHOTS["risks"][1])
    doc.add_paragraph("El modulo de riesgos registra exposicion actual, planes de respuesta y evolucion del riesgo.")
    add_table(doc, ["Campo", "Descripcion"], [
        ["Riesgo", "Descripcion del evento incierto o condicion que puede afectar el proyecto."],
        ["Probabilidad", "Valor de 1 a 5."],
        ["Impacto", "Valor de 1 a 5."],
        ["Nivel", "Calculado automaticamente como Bajo, Medio, Alto o Critico."],
        ["Estrategia", "Respuesta general: mitigar, transferir, aceptar, evitar u otra."],
        ["Mitigacion", "Acciones para reducir probabilidad o impacto."],
        ["Contingencia", "Acciones si el riesgo se materializa."],
        ["Estado", "Activo, Materializado o Cerrado."],
        ["Fecha de materializacion", "Fecha en la que el evento ocurrio, si aplica."],
        ["Impacto real / observaciones", "Efecto observado y seguimiento."],
    ])
    add_note(doc, "Score de riesgos", "El sistema evalua principalmente la exposicion actual. Un riesgo materializado o critico pesa mas que una lista larga de riesgos bajos.")


def add_conversations(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("11. Conversaciones", level=1)
    add_screenshot(doc, shots["conversations"], SCREENSHOTS["conversations"][1])
    doc.add_paragraph("Conversaciones convierte acuerdos, decisiones, bloqueos y seguimientos en informacion trazable dentro del proyecto.")
    add_table(doc, ["Funcion", "Uso"], [
        ["Nuevo hilo", "Crea una conversacion asociada al proyecto o a un contexto operativo."],
        ["Categorias", "Seguimiento, Acuerdo, Bloqueo o Decision."],
        ["Mensajes", "Registran autor, tipo y contenido."],
        ["Bloqueos", "La IA puede detectar bloqueos reportados y convertirlos en recomendaciones."],
        ["Historial", "Las decisiones, acuerdos y bloqueos relevantes alimentan trazabilidad del proyecto."],
    ])


def add_knowledge(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("12. Conocimiento, entregables y evidencias", level=1)
    add_screenshot(doc, shots["knowledge"], SCREENSHOTS["knowledge"][1])
    doc.add_paragraph("Este modulo organiza componentes, productos, entregables, evidencias cargadas e inteligencia del proyecto.")
    add_table(doc, ["Bloque", "Funciones"], [
        ["Componentes", "Registrar nombre, metodologia, responsable, objetivo y avance por componente."],
        ["Productos y evidencias", "Registrar entregables, productos de conocimiento, informes o evidencias, con estado, fecha y URL."],
        ["Evidencias cargadas", "Adjuntar archivos reales asociados a proyecto, tarea, entregable, riesgo o componente."],
        ["Descarga", "Cada evidencia cargada genera una URL de descarga controlada por permisos."],
        ["Inteligencia del proyecto", "Resume estado, riesgos detectados, hitos comprometidos y recomendaciones."],
        ["Historial de cambios", "Muestra acciones registradas por el sistema o usuarios."],
    ])


def add_ai(doc: Document, shots: dict[str, Path]) -> None:
    doc.add_heading("13. IA del Proyecto", level=1)
    add_screenshot(doc, shots["ai"], SCREENSHOTS["ai"][1])
    doc.add_paragraph("IA del Proyecto analiza la informacion disponible y produce hallazgos y recomendaciones pendientes de aprobacion humana.")
    add_table(doc, ["Elemento", "Descripcion"], [
        ["Motor activo", "Puede usar motor interno o proveedor configurado."],
        ["Analizar proyecto", "Revisa cronograma, riesgos, recursos, entregables, evidencias, presupuesto, historial y conversaciones."],
        ["Resumen ejecutivo IA", "Explica salud del proyecto y senales principales."],
        ["Hallazgos", "Lista problemas detectados por severidad y entidad relacionada."],
        ["Recomendaciones", "Acciones sugeridas con prioridad, modulo, justificacion, impacto esperado y estado."],
        ["Ver", "Abre detalle de la recomendacion y payload propuesto."],
        ["Aprobar", "Autoriza una recomendacion sin aplicarla todavia."],
        ["Rechazar", "Descarta una recomendacion."],
        ["Aplicar", "Ejecuta la accion propuesta, si el usuario tiene permisos."],
        ["Deshacer", "Revierte una recomendacion aplicada cuando la accion lo permite."],
    ])
    add_note(doc, "Control humano", "Ningun cambio recomendado por IA se aplica automaticamente. El usuario decide si aprueba, rechaza o aplica.")


def add_integrations(doc: Document) -> None:
    doc.add_heading("14. Importacion, exportacion y contratos", level=1)
    doc.add_heading("Importacion CSV", level=2)
    doc.add_paragraph("La importacion permite crear o actualizar un proyecto completo desde un archivo estructurado.")
    add_table(doc, ["Entidad CSV", "Que crea o actualiza"], [
        ["project", "Datos generales, fechas, presupuesto, metodologia y contexto estrategico."],
        ["component", "Componentes metodologicos del proyecto."],
        ["resource", "Recursos, roles, correos y capacidad."],
        ["budget", "Plan mensual por rubro con monto planificado, ejecutado y notas."],
        ["task", "Tareas, hitos, tareas resumen, WBS, responsables, avance y presupuesto."],
        ["dependency", "Relaciones de precedencia entre tareas."],
        ["sprint", "Sprints y metas."],
        ["story", "Historias Scrum, puntos, prioridad, responsable y vinculo a Plan Maestro."],
        ["risk", "Riesgos con probabilidad, impacto, respuesta y responsable."],
        ["deliverable", "Entregables, productos de conocimiento, evidencias e informes."],
        ["conversation_thread / conversation_message", "Hilos y mensajes operativos."],
    ])
    doc.add_heading("Exportaciones disponibles", level=2)
    add_bullets(doc, [
        "JSON: descarga completa del proyecto con entidades relacionadas, metricas e inteligencia.",
        "CSV: archivo compatible con la estructura de importacion.",
        "HTML: reporte ejecutivo descargable con resumen, recomendaciones, actividades, riesgos y entregables.",
    ])
    doc.add_heading("Contratos API", level=2)
    doc.add_paragraph("El backend publica OpenAPI y el proyecto genera contratos TypeScript para que frontend y API mantengan consistencia.")


def add_workflows(doc: Document) -> None:
    doc.add_heading("15. Flujos recomendados", level=1)
    add_table(doc, ["Flujo", "Pasos"], [
        ["Crear proyecto desde cero", "Portafolio > Nuevo proyecto > completar datos generales > crear > construir Plan Maestro > agregar recursos, riesgos y evidencias."],
        ["Importar proyecto completo", "Portafolio > Importar CSV > seleccionar archivo > validar entidades > abrir proyecto importado."],
        ["Gestionar cronograma", "Plan Maestro > crear tareas e hitos > indentar para WBS > vincular precedencias > recalcular > revisar ruta critica."],
        ["Planear presupuesto mensual", "Presupuesto > seleccionar mes > registrar rubro, plan y ejecutado > revisar desviacion y fuente del PHS."],
        ["Ejecutar Scrum conectado", "Scrum > crear historias > vincular a tarea del Plan Maestro > mover estados > sincronizar avance con Plan Maestro."],
        ["Gestionar riesgo materializado", "Riesgos > registrar o actualizar riesgo > estado Materializado > fecha > impacto real > contingencia."],
        ["Documentar decision", "Conversar > nuevo hilo o hilo existente > tipo Decision o Acuerdo > guardar mensaje."],
        ["Cerrar evidencia", "Conocimiento > agregar entregable > adjuntar evidencia > descargar si se requiere soporte."],
        ["Usar IA con control", "IA Proyecto > Analizar proyecto > revisar hallazgos > ver recomendacion > aprobar/rechazar/aplicar."],
    ])


def add_glossary(doc: Document) -> None:
    doc.add_heading("16. Glosario y criterios de salud", level=1)
    add_table(doc, ["Termino", "Definicion"], [
        ["PHS", "Project Health Score. Puntaje de salud del proyecto calculado con cronograma, presupuesto y riesgos."],
        ["Cronograma", "Dimension que responde si el proyecto avanza de acuerdo con fechas, hitos y actividades."],
        ["Presupuesto", "Dimension que compara presupuesto ejecutado contra presupuesto esperado a la fecha."],
        ["Riesgos", "Dimension que mide exposicion actual segun riesgos activos, criticos o materializados."],
        ["WBS", "Estructura jerarquica de trabajo. Organiza fases, actividades y tareas por niveles."],
        ["Tarea resumen", "Elemento con hijos. Su inicio, fin, duracion y avance se calculan automaticamente."],
        ["Hito", "Punto de control sin duracion."],
        ["Ruta critica", "Secuencia de tareas que impacta la fecha final del proyecto."],
        ["Riesgo materializado", "Riesgo que ocurrio y afecta actualmente el proyecto."],
    ])
    doc.add_heading("Formula PHS", level=2)
    doc.add_paragraph("PHS = Cronograma x 45% + Presupuesto x 30% + Riesgos x 25%.")
    add_table(doc, ["Estado", "Criterio operativo"], [
        ["Saludable", "PHS mayor o igual a 80."],
        ["En riesgo", "PHS menor a 80 y mayor o igual a 60."],
        ["Critico", "PHS menor a 60."],
    ])


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    shots = crop_images()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_cover(doc)
    add_toc(doc)
    add_overview(doc, shots)
    add_auth(doc, shots)
    add_common(doc)
    add_portfolio(doc, shots)
    add_project_create(doc)
    add_master(doc, shots)
    add_scrum(doc, shots)
    add_resources(doc, shots)
    add_budget(doc, shots)
    add_risks(doc, shots)
    add_conversations(doc, shots)
    add_knowledge(doc, shots)
    add_ai(doc, shots)
    add_integrations(doc)
    add_workflows(doc)
    add_glossary(doc)

    for i, section in enumerate(doc.sections):
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Manual de Usuario Proyecta360"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
